from flask import Flask, request, render_template, send_file, send_from_directory
import pandas as pd
import docx
import os


def crear_app():
    app = Flask(__name__)

    # Definimos la ruta donde se guardarán los archivos subidos
    CARPETA_CARGA = 'uploads'

    # Creamos una carpeta de uploads si no existe
    if not os.path.exists(CARPETA_CARGA):
        os.makedirs(CARPETA_CARGA)

    # Ruta para la cual mandamos a llamar al archivo HTML
    @app.route('/')
    def index():
        return render_template('index.html')

    # Ruta para subir archivos, solo solicitudes POST
    @app.route('/upload', methods=['POST'])
    def subir_archivos():
        if 'file' not in request.files:
            mensaje_error = 'No se ha cargado ningún archivo.'
            tipo_mensaje = 'danger'
            return render_template('index.html', mensaje=mensaje_error, categoria=tipo_mensaje)

        archivo = request.files['file']

        if archivo.filename == '':
            mensaje_error = 'Archivo no seleccionado.'
            tipo_mensaje = 'danger'
            return render_template('index.html', mensaje=mensaje_error, categoria=tipo_mensaje)

        if not archivo.filename.lower().endswith('.docx'):
            mensaje_error = 'Archivo no compatible por el momento, por favor cargue un archivo .docx.'
            tipo_mensaje = 'danger'
            return render_template('index.html', mensaje=mensaje_error, categoria=tipo_mensaje)

        if archivo:
            ruta_archivo = os.path.join(CARPETA_CARGA, 'input.docx')
            archivo.save(ruta_archivo)

            ruta_salida = os.path.join(CARPETA_CARGA, 'Reporte.csv')
            try:
                convert_docx_to_csv(ruta_archivo, ruta_salida)
            except Exception as e:
                mensaje_error = f'Error al convertir el archivo: {e}'
                tipo_mensaje = 'danger'
                return render_template('index.html', mensaje=mensaje_error, categoria=tipo_mensaje)

            return render_template('index.html', mensaje='¡Archivo convertido con éxito!', descargar=True, archivo=ruta_salida)

        return render_template('index.html', mensaje='No se ha seleccionado un archivo.', categoria='danger')

    # Ruta para descargar el archivo convertido
    @app.route('/descargar/<path:archivo>')
    def descargar(archivo):
        return send_file(archivo, as_attachment=True)
    
    # Ruta para servir archivos estáticos en el directorio 'assets'
    @app.route('/assets/<path:path>')
    def send_asset(path):
        return send_from_directory('assets', path)

    def addReporte(df, movimientos, solicitud):
        for reg in range(1, len(movimientos)):
            df.loc[len(df)] = {
                'ORDEN': solicitud,
                'TIEMPO DE INICIO': movimientos[reg][1] + " " + movimientos[reg][2],
                'ACTIVIDAD': movimientos[reg][0],
                'RESPONSABLE': movimientos[reg][4],
                'OBSERVACIONES': movimientos[reg][6]
            }
        return df

    def convert_docx_to_csv(input_path, output_path):
        archivo = docx.Document(input_path)
        tablas = archivo.tables
        parrafos = archivo.paragraphs

        reporte = pd.DataFrame(columns=['ORDEN', 'TIEMPO DE INICIO', 'ACTIVIDAD', 'RESPONSABLE', 'OBSERVACIONES'])
        indP = indT = indB = 0
        indTmovimientos = 0
        solicitud = ''
        for parrafo in archivo.paragraphs:
            if parrafo.text == '':
                indP += 1
                continue
            if parrafo.text == 'INFORMACIÓN GENERAL':
                solicitud = [[cell.text for cell in row.cells] for row in tablas[indT].rows][0][1]
            if parrafo.text == 'BITACORA DE MOVIMIENTOS':
                indTmovimientos += 1
                movimientos = [[cell.text for cell in row.cells] for row in tablas[indB].rows]
                reporte = addReporte(reporte, movimientos, solicitud)
            indP += 1
            indT += 1
            indB += 1
        print('Se encontraron ' + str(indTmovimientos) + ' tablas de movimientos.')
        reporte.to_csv(output_path, index=False)

    return app


if __name__ == '__main__':
    app = crear_app()
    app.run()
